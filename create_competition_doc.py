# -*- coding: utf-8 -*-
"""
Create the competition document for 2026中国高校计算机大赛人工智能创意赛（鸿蒙赛道）
Based on the template structure, filling in all content sections.
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_font(cell, text, font_name='宋体', font_size=Pt(10.5), bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Set cell text with specific font settings."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold

def add_formatted_paragraph(doc, text, font_name='宋体', font_size=Pt(10.5), bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(0), space_before=Pt(0), first_line_indent=None):
    """Add a paragraph with specific formatting."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold
    return p

def add_heading_custom(doc, text, level=1):
    """Add heading with custom font setting."""
    font_sizes = {0: Pt(22), 1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(10.5)}
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if level in font_sizes:
            run.font.size = font_sizes[level]
    return h

def create_document():
    doc = Document()

    # ========== Page Setup ==========
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    # ========== COVER PAGE ==========
    # Title
    add_formatted_paragraph(doc, '', font_size=Pt(10.5))
    add_formatted_paragraph(doc, '', font_size=Pt(10.5))
    add_formatted_paragraph(doc, '2026中国高校计算机大赛—人工智能创意赛',
                           font_size=Pt(22), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_after=Pt(20))

    add_formatted_paragraph(doc, '', font_size=Pt(10.5))
    add_formatted_paragraph(doc, '鸿蒙赛道作品说明文档（初赛）',
                           font_size=Pt(18), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_after=Pt(30))

    add_formatted_paragraph(doc, '', font_size=Pt(10.5))
    add_formatted_paragraph(doc, '', font_size=Pt(10.5))

    # Cover info
    cover_items = [
        ('参赛学校：', '内蒙古民族大学'),
        ('团队名称：', '伏特加代码'),
        ('作品名称：', '灵宠家伴'),
        ('赛题方向：', 'Agent创新'),
        ('联系人：', '杨彦浩'),
        ('联系电话：', '15047319273'),
    ]
    for label, value in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(label + '_______' + value + '__________')
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(14)

    doc.add_page_break()

    # ========== PAGE 2: TEAM INFO TABLE ==========
    add_formatted_paragraph(doc, '2026中国高校计算机大赛—人工智能创意赛',
                           font_size=Pt(14), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_after=Pt(6))
    add_formatted_paragraph(doc, '参赛团队信息表',
                           font_size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_after=Pt(12))

    # Basic info table
    table1 = doc.add_table(rows=4, cols=2, style='Table Grid')
    table1.autofit = True

    info_data = [
        ('作品名称', '灵宠家伴'),
        ('团队名称', '伏特加代码'),
        ('参赛学校', '内蒙古民族大学'),
        ('参赛赛题方向', '（ ）应用创新  （ √ ）Agent创新  （ ）用户体验创新  （ ）操作系统智能创新'),
    ]
    for i, (label, value) in enumerate(info_data):
        set_cell_font(table1.cell(i, 0), label, font_size=Pt(10.5), bold=True)
        set_cell_font(table1.cell(i, 1), value, font_size=Pt(10.5), alignment=WD_ALIGN_PARAGRAPH.LEFT)
        table1.cell(i, 0).width = Cm(3)

    add_formatted_paragraph(doc, '', font_size=Pt(6))
    add_formatted_paragraph(doc, '团队队员基本信息', font_size=Pt(12), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    # Team member table
    headers = ['姓名', '学校全称', '院（系）全称', '专业全称', '年级', '毕业时间', '联系电话', '邮箱', '团队分工']
    members = [
        ['杨彦浩', '内蒙古民族大学', '计算机科学与技术学院', '软件工程', '大三', '2027.6', '15047319273', 'yangyanhao2004@qq.com', '队长'],
        ['倪文博', '内蒙古民族大学', '计算机科学与技术学院', '计算机科学与技术', '大三', '2027.6', '15645187267', '2775112791@qq.com', '队员'],
        ['林炜彦', '内蒙古民族大学', '计算机科学与技术学院', '网络工程', '大三', '2027.3', '15302929025', '1484351341@qq.com', '队员'],
    ]

    table2 = doc.add_table(rows=len(members)+1, cols=len(headers), style='Table Grid')
    table2.autofit = True

    for j, h in enumerate(headers):
        set_cell_font(table2.cell(0, j), h, font_size=Pt(8), bold=True)
    for i, row in enumerate(members):
        for j, val in enumerate(row):
            set_cell_font(table2.cell(i+1, j), val, font_size=Pt(8))

    add_formatted_paragraph(doc, '', font_size=Pt(6))
    add_formatted_paragraph(doc, '团队指导教师信息（指导老师须与队长同校）',
                           font_size=Pt(12), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    # Teacher table
    teacher_headers = ['姓名', '院（系）全称', '职称', '研究方向', '联系电话', '联系邮箱']
    teacher_row = ['路扬', '计算机科学与技术学院', '教授', '深度学习、自然语言处理', '15004997746', '']

    table3 = doc.add_table(rows=2, cols=len(teacher_headers), style='Table Grid')
    table3.autofit = True
    for j, h in enumerate(teacher_headers):
        set_cell_font(table3.cell(0, j), h, font_size=Pt(9), bold=True)
    for j, val in enumerate(teacher_row):
        set_cell_font(table3.cell(1, j), val, font_size=Pt(9))

    add_formatted_paragraph(doc, '', font_size=Pt(6))
    add_formatted_paragraph(doc, '团队成员优势描述',
                           font_size=Pt(12), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    advantages = (
        '杨彦浩：负责系统总体架构设计与嵌入式开发（STM32 + FreeRTOS），精通C/Python/ArkTS多语言开发，'
        '具备从底层硬件驱动到上层应用的全栈开发能力；\n'
        '倪文博：负责系统后端架构设计与Java Spring Boot数据服务开发，熟悉WebSocket实时通信与JWT认证体系，'
        '负责AI Agent功能集成与Kimi API调用链设计；\n'
        '林炜彦：负责系统测试、安全审计与漏洞挖掘，熟悉WebRTC实时通信协议，'
        '负责多端协同（Web + 鸿蒙）的联调测试与性能优化。'
    )
    add_formatted_paragraph(doc, advantages, font_size=Pt(10.5))

    doc.add_page_break()

    # ========== PAGE 3: ORIGINALITY DECLARATION ==========
    add_formatted_paragraph(doc, '灵宠家伴     作品原创性声明',
                           font_size=Pt(18), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_after=Pt(20))

    declaration = (
        '郑重声明：承诺本参赛队伍报名信息真实有效；呈交的参赛作品相关资料以及所完成的作品实物等相关成果，'
        '是本团队独立进行研究工作所取得的成果，除文中已经注明引用的内容外，本作品说明文档不包含任何其他个人或集体'
        '已经发表或撰写过的作品成果，不侵犯任何第三方的知识产权或其他权利。本声明的法律结果由本参赛队承担。'
    )
    add_formatted_paragraph(doc, declaration, font_size=Pt(12), first_line_indent=Cm(0.74), space_after=Pt(30))

    add_formatted_paragraph(doc, '参赛队员签名（团队全部成员）：', font_size=Pt(12), space_after=Pt(10))
    add_formatted_paragraph(doc, '杨彦浩          倪文博          林炜彦', font_size=Pt(12), space_after=Pt(20))
    add_formatted_paragraph(doc, '日期： 2026   年   6   月  10   日', font_size=Pt(12), space_after=Pt(30))

    add_formatted_paragraph(doc, '指导老师审核签名：', font_size=Pt(12), space_after=Pt(10))
    add_formatted_paragraph(doc, '', font_size=Pt(12))
    add_formatted_paragraph(doc, '日期：    年      月     日', font_size=Pt(12), space_after=Pt(20))

    note = (
        '（注：本页签名可以打印纸质文件后签名，提供扫描件；或者使用电子签名。'
        '不论哪种方式，需保证页面内容完整、字迹清晰，请与本项目创意书作为同一文档提交，'
        '否则项目创意书提交可能无效。）'
    )
    add_formatted_paragraph(doc, note, font_size=Pt(9))

    doc.add_page_break()

    # ========== PAGE 4: SUBMISSION GUIDELINES ==========
    add_formatted_paragraph(doc, '作品说明文档提交规范说明',
                           font_size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_after=Pt(10))

    guidelines = (
        '将综合考察评估各参赛队的创新性、完备性、前景评估、规范性、实际应用价值及团队成员背景及能力互补。'
        '因此，为了更加清晰了解各参赛队的基本情况及作品内容，各参赛队（请以团队为单位）提交的《作品说明文档》'
        '将用于评委评审打分的核心依据。'
    )
    add_formatted_paragraph(doc, guidelines, font_size=Pt(10.5), first_line_indent=Cm(0.74), space_after=Pt(6))

    add_formatted_paragraph(doc, '作品说明文档，具体内容：', font_size=Pt(10.5), bold=True, space_after=Pt(6))

    sections_guide = [
        ('一、创意描述', '一句话描述关键创新点，不超过30字。'),
        ('二、设计稿（应用/Agent/用户体验）/技术方案（操作系统智能）',
         '提供创意效果图，或交互流程图，展示创意的视觉设计与交互逻辑；鼓励应用、Agent、用户体验三个赛题输出宣传海报。'),
        ('三、介绍文档',
         '具体描述创意的内容和实现路径，可包含创意背景（行业痛点与用户需求分析）、核心功能设计（交互逻辑与技术实现路径）、及市场前景评估。不超过800字，"操作系统智能赛题"需要额外提供测试报告。'),
    ]
    for title, desc in sections_guide:
        add_formatted_paragraph(doc, f'{title}', font_size=Pt(10.5), bold=True, space_after=Pt(2))
        add_formatted_paragraph(doc, desc, font_size=Pt(10.5), first_line_indent=Cm(0.74), space_after=Pt(6))

    add_formatted_paragraph(doc, '二、格式及其他要求', font_size=Pt(12), bold=True, space_after=Pt(4))
    format_items = [
        '（一）字体：宋体',
        '（二）字号：标题二号粗体，一级标题三号粗体，二级标题四号粗体，三级标题小四号粗体，正文五号',
        '（三）行距：单倍行距',
        '（四）页面设置：上2.5cm、下2.5cm、左3cm、右3cm、装订线0cm；页眉1.5cm、页脚1.5cm；A4纵向',
        '（五）插入页码：页面底端，对齐方式外侧',
        '（六）其他：PDF格式提交，命名格式为 01-作品说明文档+参赛队伍名称；文档中若包含插图、表格请按序编号并命名；主体内容不超过20页',
    ]
    for item in format_items:
        add_formatted_paragraph(doc, item, font_size=Pt(9))

    doc.add_page_break()

    # ====================================================================
    # MAIN CONTENT STARTS HERE
    # ====================================================================

    # ========== SECTION 1: 创意描述 ==========
    add_heading_custom(doc, '一、创意描述', level=1)
    add_formatted_paragraph(doc, '', font_size=Pt(6))

    innovation_text = 'AI Agent理解自然语言，能直接操控物理机器人执行运动指令。'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(innovation_text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14)
    run.bold = True

    doc.add_page_break()

    # ========== SECTION 2: 设计稿/技术方案 ==========
    add_heading_custom(doc, '二、设计稿（Agent创新——技术方案与交互设计）', level=1)

    # 2.1 System Architecture
    add_heading_custom(doc, '2.1 系统总体架构', level=2)

    arch_text = (
        '灵宠家伴系统采用"四子系统+多端协同"的分层架构设计，从底层嵌入式硬件到上层应用全链路打通。'
        '系统由树莓派实时运行时（Python FastAPI，端口8080）、Web前端控制面板（Vue 3 + TypeScript）、'
        'Java Spring Boot数据后端（端口8090）、STM32嵌入式运动控制固件（C + FreeRTOS）、'
        '以及鸿蒙原生移动端App（ArkTS）五大模块组成。'
    )
    add_formatted_paragraph(doc, arch_text, font_size=Pt(10.5), first_line_indent=Cm(0.74), space_after=Pt(6))

    # Architecture layers table
    add_formatted_paragraph(doc, '表2.1 系统架构分层表', font_size=Pt(9), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

    arch_table = doc.add_table(rows=5, cols=3, style='Table Grid')
    arch_table.autofit = True
    arch_headers = ['层次', '技术栈', '核心职责']
    arch_data = [
        ['表示层', 'Vue 3 + TypeScript + Tailwind CSS\n鸿蒙 ArkTS 原生UI', 'Web与鸿蒙双端用户界面，包括控制面板、AI对话、相册管理、提醒管理、用户认证'],
        ['实时服务层', 'Python FastAPI + WebSocket\nWebRTC (aiortc) + MJPEG', 'WebSocket消息中继路由、MJPEG视频流分发、WebRTC信令转发、AI对话处理与指令解析'],
        ['数据服务层', 'Java Spring Boot 3 + JPA\nMySQL 8.0 + JWT + SMTP', '用户认证授权、提醒CRUD与定时调度、相册管理、操作日志记录、邮件发送'],
        ['物理嵌入式层', 'STM32F407 + FreeRTOS\nPID控制 + 编码器 + PWM', '直流电机闭环PID速度控制、舵机角度控制、编码器速度反馈、UART指令解析'],
    ]
    for j, h in enumerate(arch_headers):
        set_cell_font(arch_table.cell(0, j), h, font_size=Pt(9), bold=True)
    for i, row in enumerate(arch_data):
        for j, val in enumerate(row):
            set_cell_font(arch_table.cell(i+1, j), val, font_size=Pt(8), alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_formatted_paragraph(doc, '', font_size=Pt(4))

    # 2.2 Agent Design
    add_heading_custom(doc, '2.2 AI Agent智能体设计', level=2)

    agent_text = (
        '本项目的核心创新在于构建了一个"具身智能Agent"（Embodied AI Agent），将大语言模型的推理能力'
        '与物理机器人的执行能力深度结合。Agent通过以下机制实现自然语言到物理动作的转化：'
    )
    add_formatted_paragraph(doc, agent_text, font_size=Pt(10.5), first_line_indent=Cm(0.74), space_after=Pt(6))

    # Agent workflow
    add_formatted_paragraph(doc, '表2.2 AI Agent功能调用体系', font_size=Pt(9), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

    agent_table = doc.add_table(rows=8, cols=3, style='Table Grid')
    agent_table.autofit = True
    agent_headers = ['工具名称', '触发方式', '功能描述']
    agent_data = [
        ['control_robot', '用户说出运动指令\n（如"前进5秒"）', 'AI调用function calling，通过WebSocket→Pi→UART→STM32链\r\n路控制机器人执行前进/后退/左转/右转/停止等动作'],
        ['set_reminder', '用户说出提醒需求\n（如"3分钟后提醒我"）', 'AI解析提醒内容和时间，通过HTTP POST调用Java后端创建提醒，\r\n到时间后通过EMAIL或VOICE（TTS语音播报）方式送达'],
        ['chat_reply', '通用对话场景', 'AI直接返回文字回复，支持多轮对话上下文（每用户最多50轮）'],
        ['time', '关键词"时间/几点"', '返回当前系统日期和时间'],
        ['weather', '关键词"天气"', '通过OpenWeatherMap API查询城市天气'],
        ['news', '关键词"新闻"', '通过NewsAPI获取最新头条新闻'],
        ['joke', '关键词"笑话"', '通过Official Joke API获取随机笑话'],
    ]
    for j, h in enumerate(agent_headers):
        set_cell_font(agent_table.cell(0, j), h, font_size=Pt(9), bold=True)
    for i, row in enumerate(agent_data):
        for j, val in enumerate(row):
            set_cell_font(agent_table.cell(i+1, j), val, font_size=Pt(8), alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_formatted_paragraph(doc, '', font_size=Pt(4))

    # 2.3 Intent recognition enhancement
    add_heading_custom(doc, '2.3 意图识别增强方案（Scheme A+B）', level=2)

    scheme_text = (
        '为确保操控类指令的响应实时性和准确性，系统设计了双方案组合的意图识别增强机制：\n'
        '• Scheme A — 预过滤层：使用正则匹配直接拦截"前进/后退/左转/右转/停下"等清晰运动指令，'
        '命中后跳过AI调用，直接通过WebSocket发送控制指令到Pi端，延迟≤50ms；\n'
        '• Scheme B — 强制tool_choice：对于检测到移动类/提醒类关键词的用户输入，'
        '调用Kimi K2.5 API时设置force tool_choice参数，确保AI调用control_robot或set_reminder工具；\n'
        '• 经过A+B方案增强后，常见操控指令的响应延迟从LLM调用的3-5秒降低至50ms以内，'
        '识别准确率从约70%提升至95%以上。'
    )
    add_formatted_paragraph(doc, scheme_text, font_size=Pt(10.5), first_line_indent=Cm(0.74))

    add_formatted_paragraph(doc, '', font_size=Pt(4))

    # 2.4 HarmonyOS App Design
    add_heading_custom(doc, '2.4 鸿蒙端应用设计', level=2)

    harmony_text = (
        '鸿蒙原生App（Family_Robot_HarmonyOS）是系统在HarmonyOS生态中的核心入口，'
        '采用ArkTS语言开发，适配华为MatePad Air等设备。App完整实现了Web端的全部核心功能：\n'
        '• 实时控制面板：方向键（D-Pad）控制机器人移动、云台（PTZ）舵机角度调节、三档速度切换；\n'
        '• 视频监控：通过HTTP轮询/video/frame端点获取实时JPEG快照（鸿蒙不支持MJPEG流，采用帧轮询方案）；\n'
        '• AI对话：通过WebSocket发送ai_chat消息与Kimi K2.5 Agent交互，支持机器人控制与提醒创建；\n'
        '• 用户体系：JWT认证登录/注册/密码重置，与Java后端REST API对接；\n'
        '• 相册管理：照片网格展示、批量下载/删除、照片上传；\n'
        '• 提醒管理：提醒CRUD、启用/禁用切换、按时间查看。\n\n'
        '鸿蒙端的WebSocket服务完全对齐Vue前端协议，实现了跨平台统一通信标准。'
    )
    add_formatted_paragraph(doc, harmony_text, font_size=Pt(10.5), first_line_indent=Cm(0.74))

    add_formatted_paragraph(doc, '', font_size=Pt(4))

    # 2.5 Communication protocol design
    add_heading_custom(doc, '2.5 通信协议设计', level=2)

    comm_text = (
        '系统设计了四层通信协议体系，实现多子系统间的实时数据交换：\n'
        '（1）WebSocket实时协议（端口8080）：Web/鸿蒙↔Python↔Pi之间传输JSON格式的控制指令(7种)、'
        '机器人状态(battery/signal/motor_speed)、AI对话消息(ai_chat/chat_reply)、'
        'WebRTC信令(offer/answer/candidate)、摄像头帧(base64 JPEG)；\n'
        '（2）HTTP REST协议（端口8090）：Web/鸿蒙↔Java之间传输认证信息、CRUD操作、JWT令牌；\n'
        '（3）内部HTTP调用：Python↔Java之间传输AI创建的提醒和语音提醒触发指令；\n'
        '（4）UART串口协议（115200bps）：Pi↔STM32之间传输V1/S1/M1格式的电机/舵机指令和编码器反馈。'
    )
    add_formatted_paragraph(doc, comm_text, font_size=Pt(10.5), first_line_indent=Cm(0.74))

    add_formatted_paragraph(doc, '', font_size=Pt(4))

    # 2.6 UI Design Description
    add_heading_custom(doc, '2.6 交互界面设计', level=2)

    ui_text = (
        '灵宠家伴系统的用户界面设计遵循"简洁直观、功能优先"的原则，提供Web和鸿蒙双端一致的操作体验：\n\n'
        '（1）主控面板（Dashboard）：左侧为实时视频画面区域，展示MJPEG摄像头画面；右侧为控制区域，'
        '包含十字方向键（D-Pad）控制机器人移动、WASD云台控制区调节摄像头角度、速度档位选择器。'
        '顶部状态栏展示WebSocket连接状态、电池电量百分比、WiFi信号强度等关键信息。\n\n'
        '（2）AI对话界面（AI Chat）：类聊天应用的消息气泡布局，用户输入自然语言指令后，'
        'Agent返回带有执行动作标记的回复消息（如"[正在前进2秒]"），提供清晰的操作确认反馈。'
        '支持多轮对话，显示完整对话历史。\n\n'
        '（3）认证页面：采用卡片式布局的登录/注册/密码重置页面，支持邮箱验证码和密码两种登录方式。\n\n'
        '（4）管理页面：相册采用网格布局展示照片缩略图，支持多选批量操作；提醒管理采用列表布局，'
        '显示提醒内容、时间、方式（邮件/语音）、启用状态，支持一键开关。\n\n'
        '（5）鸿蒙端适配：鸿蒙App针对平板设备进行了大屏优化，控制面板采用左右分栏布局，'
        '视频区域与控制按钮分区展示，充分利用HarmonyOS的响应式布局能力。'
        '所有页面组件（StatusBar、VideoStream、DPadControl、PTZControl、ChatBubble等）'
        '均封装为独立的可复用[@Component]装饰器组件。'
    )
    add_formatted_paragraph(doc, ui_text, font_size=Pt(10.5), first_line_indent=Cm(0.74))

    doc.add_page_break()

    # ========== SECTION 3: 介绍文档 ==========
    add_heading_custom(doc, '三、介绍文档', level=1)

    # 3.1 Background
    add_heading_custom(doc, '3.1 创意背景（行业痛点与用户需求分析）', level=2)

    background_text = (
        '随着物联网和AI技术的快速发展，智能家居已从简单的设备联动升级为具备感知、决策与执行能力的综合服务系统。'
        '然而，当前市场上面向家庭的智能机器人产品普遍存在以下痛点：\n\n'
        '第一，功能碎片化严重：远程操控、视频监控、语音助手、日程提醒等功能分散在不同设备和应用中，'
        '缺乏统一的操作入口和协同机制，用户需频繁切换多个App或设备。\n\n'
        '第二，AI与物理设备脱节：多数AI语音助手（如智能音箱）仅能提供信息查询和媒体播放，'
        '无法直接控制物理设备（如机器人移动、云台摄像头转动等）。用户"说"出来的指令无法驱动"实体"执行，'
        '"理解"与"行动"之间存在断裂。\n\n'
        '第三，实时性不足：市面上的远程操控方案（如网络摄像头+遥控车）普遍存在视频延迟大、控制响应慢的问题，'
        '无法实现真正的"临场感"远程操控体验。\n\n'
        '第四，定制化门槛高：商用机器人系统封闭，API不开放或仅提供有限的云端对接能力，'
        '开发者和个人用户难以根据实际需求进行功能扩展和系统定制。\n\n'
        '第五，鸿蒙生态缺失：在HarmonyOS生态中，目前缺少面向家庭服务机器人场景的原生端侧应用，'
        '无法充分发挥鸿蒙分布式能力和跨设备协同优势。'
    )
    add_formatted_paragraph(doc, background_text, font_size=Pt(10.5), first_line_indent=Cm(0.74))

    # 3.2 Core Features
    add_heading_custom(doc, '3.2 核心功能设计与技术实现路径', level=2)

    core_text = (
        '灵宠家伴系统的核心创新在于构建了一个"具身智能Agent"——将大语言模型（Kimi K2.5）的推理决策能力'
        '与物理机器人（STM32控制的轮式底盘）的执行能力通过多层协议打通，实现"说话即操控"的自然交互体验。'
        '具体实现路径如下：\n\n'
        '（一）多模态交互通道\n'
        '系统提供三条并行的交互通道：Web浏览器（Vue 3 SPA）、鸿蒙原生App（ArkTS）、以及树莓派端本地语音'
        '（openWakeWord唤醒词"Hey Jarvis"→Whisper.cpp语音识别→Router路由→Kimi LLM→Piper TTS语音合成）。'
        '三条通道通过同一套WebSocket协议和REST API接口接入系统，实现跨端统一体验。\n\n'
        '（二）AI Agent智能决策引擎\n'
        'Agent基于Kimi K2.5大语言模型，通过function calling机制调用7个工具函数（control_robot、'
        'set_reminder、chat_reply及4个信息查询工具）。Agent的创新之处在于：\n'
        '（1）双方案意图识别增强——预过滤正则匹配（Scheme A）直接拦截清晰指令，绕过AI调用，延迟≤50ms；'
        '强制tool_choice机制（Scheme B）确保模糊指令也能准确路由到控制工具。\n'
        '（2）命令队列机制——当用户一口气说出多个指令时（如"前进5秒然后左转再拍照"），'
        'AI解析后按序加入命令队列，通过duration等待确保每个动作完整执行后再执行下一个，避免指令冲突。\n'
        '（3）会话上下文管理——每用户维护最多50轮对话历史，支持多轮指代消解（如"再快一点"关联到上一轮的速度指令）。\n\n'
        '（三）实时通信与硬件控制链路\n'
        '控制指令从用户发出到硬件执行的完整链路为：用户输入→WebSocket→Python FastAPI→WebSocket→'
        'Pi Client→UART 115200bps→STM32→PID闭环控制→电机PWM输出。在局域网环境下，'
        '端到端控制延迟≤100ms，视频延迟≤500ms。STM32固件运行FreeRTOS实时操作系统，'
        '10ms级PID控制中断保证速度闭环的控制精度。\n\n'
        '（四）智能提醒服务\n'
        'Agent支持通过对话自动创建提醒（如"5分钟后提醒我开会"），提醒支持EMAIL（QQ SMTP邮件）'
        '和VOICE（机器人TTS语音播报）两种方式。Java后端的每分钟定时任务（@Scheduled cron表达式）'
        '扫描到期提醒并触发发送。中文语音提醒在存储时由Kimi自动翻译为英文（Piper TTS为英文引擎），'
        '确保语音播报质量。\n\n'
        '（五）视频监控与语音通话\n'
        '系统通过Picamera2采集摄像头画面（640×360@10fps JPEG），经WebSocket→VideoStreamHub内存缓冲→'
        'HTTP MJPEG流式传输到浏览器。同时支持WebRTC（aiortc）建立P2P音频直连，'
        '实现浏览器与机器人之间的双向语音通话，延迟极低。\n\n'
        '（六）鸿蒙原生适配\n'
        '鸿蒙端App完整移植了Web端的功能体系，采用ArkTS声明式UI开发，6个页面（控制面板、AI对话、'
        '相册、提醒、设置、认证）配合6个可复用组件实现。针对鸿蒙不支持MJPEG流的限制，'
        '采用帧轮询/video/frame端点方案实现视频预览。WebSocket服务包含心跳保活（30秒）、'
        '自动重连（最多5次、1.5秒间隔）等完整的连接管理逻辑。'
    )
    add_formatted_paragraph(doc, core_text, font_size=Pt(10.5), first_line_indent=Cm(0.74))

    # 3.3 Market Prospects
    add_heading_custom(doc, '3.3 市场前景评估', level=2)

    market_text = (
        '灵宠家伴系统具备广阔的市场应用前景和良好的社会经济效益：\n\n'
        '（一）居家养老辅助\n'
        '随着中国老龄化程度加深，独居老人数量持续增长。灵宠家伴可为独居老人提供：远程子女操控探望（视频+语音）、'
        'AI日常陪伴聊天、定时用药提醒（VOICE语音播报）、异常情况远程报警等功能。'
        '相较于商用养老机器人动辄万元的售价，本系统硬件成本仅约590-930元（树莓派+STM32+摄像头+电机底盘），'
        '具有极高的性价比优势。\n\n'
        '（二）儿童陪伴教育\n'
        'AI语音助手可进行知识问答和语言学习互动，机器人运动能力增加了物理交互的趣味性。'
        '家长可通过手机远程查看家中情况，与孩子进行WebRTC实时通话，实现"远程陪娃"场景。\n\n'
        '（三）家庭安防监控\n'
        '结合可移动的机器人底盘与摄像头云台，用户可远程巡逻家中的各个角落，弥补固定摄像头的监控盲区。'
        '拍照功能可随时记录关键画面并保存到相册，形成完整的家庭安防记录。\n\n'
        '（四）开源生态价值\n'
        '项目基于全开源技术栈构建（Python + Vue + Java + FreeRTOS + 鸿蒙ArkTS），'
        '是少有的从底层硬件到上层应用全链路打通的完整开源机器人项目。可作为高校嵌入式、Web开发、'
        'AI集成、鸿蒙应用开发等课程的教学案例，具有显著的技术教育价值。\n\n'
        '（五）商业化扩展\n'
        '模块化的系统架构支持灵活扩展：可接入计算机视觉（OpenCV）实现自动巡逻与异常行为检测；'
        '可集成SLAM（如Cartographer）实现自主导航与建图；可对接HomeAssistant等智能家居平台实现设备联动；'
        '可扩展多机器人协同编队与任务分配。算力允许的条件下，可将云端Kimi替换为本地部署的大模型'
        '（如Qwen2.5/Ollama），实现完全离线的隐私保护方案。'
    )
    add_formatted_paragraph(doc, market_text, font_size=Pt(10.5), first_line_indent=Cm(0.74))

    doc.add_page_break()

    # ========== APPENDIX: Technical Details ==========
    add_heading_custom(doc, '附录A：系统技术指标', level=1)

    add_formatted_paragraph(doc, '表A.1 系统性能指标', font_size=Pt(9), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

    perf_table = doc.add_table(rows=9, cols=2, style='Table Grid')
    perf_table.autofit = True
    perf_data = [
        ['指标', '数值'],
        ['控制指令端到端延迟（局域网）', '≤ 100ms'],
        ['MJPEG视频流延迟', '≤ 500ms'],
        ['AI对话回复延迟（预过滤指令）', '≤ 50ms'],
        ['AI对话回复延迟（Kimi LLM调用）', '≤ 5s'],
        ['WebRTC音频通话延迟（P2P）', '≤ 200ms'],
        ['PID控制周期', '10ms'],
        ['摄像头帧率', '10fps @ 640×360'],
        ['STM32主频', '168MHz (Cortex-M4)'],
    ]
    for i, row in enumerate(perf_data):
        for j, val in enumerate(row):
            is_bold = (i == 0)
            set_cell_font(perf_table.cell(i, j), val, font_size=Pt(9), bold=is_bold)

    add_formatted_paragraph(doc, '', font_size=Pt(6))

    add_formatted_paragraph(doc, '表A.2 技术栈一览', font_size=Pt(9), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

    tech_table = doc.add_table(rows=6, cols=4, style='Table Grid')
    tech_table.autofit = True
    tech_headers = ['子系统', '语言/框架', '关键依赖', '部署位置']
    tech_data = [
        ['Pi实时运行时', 'Python 3 + FastAPI', 'aiortc, openWakeWord, Whisper.cpp, Piper TTS (ONNX), Picamera2, httpx', '树莓派4B (2GB+ RAM)'],
        ['Web前端', 'Vue 3 + TypeScript + Vite', 'Pinia, Tailwind CSS v4, 原生WebSocket/WebRTC API', 'PC端浏览器'],
        ['鸿蒙App', 'ArkTS (HarmonyOS 4.2+)', '@ohos.net.webSocket, router, 原生HTTP', '华为MatePad Air等设备'],
        ['Java数据后端', 'Spring Boot 3 + JPA', 'Spring Security, JWT, Java Mail, Hibernate, MySQL 8.0', 'PC端 (端口8090)'],
        ['STM32固件', 'C + FreeRTOS + HAL', 'TIM1/TIM2/TIM5/TIM7/TIM13, USART1, PID控制器', 'STM32F407VET6'],
    ]
    for j, h in enumerate(tech_headers):
        set_cell_font(tech_table.cell(0, j), h, font_size=Pt(9), bold=True)
    for i, row in enumerate(tech_data):
        for j, val in enumerate(row):
            set_cell_font(tech_table.cell(i+1, j), val, font_size=Pt(8), alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_formatted_paragraph(doc, '', font_size=Pt(6))

    add_formatted_paragraph(doc, '表A.3 AI能力矩阵', font_size=Pt(9), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

    ai_table = doc.add_table(rows=8, cols=3, style='Table Grid')
    ai_table.autofit = True
    ai_headers = ['AI能力', '技术方案', '部署位置']
    ai_data = [
        ['唤醒词检测', 'openWakeWord (hey_jarvis模型)', '树莓派本地'],
        ['语音识别(STT)', 'whisper.cpp (ggml-tiny.en-q5_0)', '树莓派本地'],
        ['语音合成(TTS)', 'Piper TTS (en_GB-semaine-medium, ONNX)', '树莓派本地'],
        ['大语言模型', 'Kimi K2.5 API (moonshot-v1-8k)', '云端 (Moonshot)'],
        ['本地LLM(可选)', 'Ollama + qwen2.5:1.5b', '树莓派本地(默认禁用)'],
        ['中文→English翻译', 'Kimi K2.5 API', '云端（提醒翻译）'],
        ['意图识别增强', '正则预过滤 + force tool_choice', 'Python Backend'],
    ]
    for j, h in enumerate(ai_headers):
        set_cell_font(ai_table.cell(0, j), h, font_size=Pt(9), bold=True)
    for i, row in enumerate(ai_data):
        for j, val in enumerate(row):
            set_cell_font(ai_table.cell(i+1, j), val, font_size=Pt(8), alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_formatted_paragraph(doc, '', font_size=Pt(8))

    add_heading_custom(doc, '附录B：关键代码与技术架构', level=1)

    tech_detail = (
        'B.1 Agent工具调用示例（control_robot）\n'
        '当用户通过Web/鸿蒙的AI对话界面输入"前进5秒"时：\n'
        '1. prefilter_command()正则匹配"前进"→命中Scheme A→直接执行，延迟≤50ms\n'
        '2. 若更复杂的指令"前进3秒后左转然后提醒我关火"→进入Scheme B\n'
        '3. _detect_tool_choice()检测到"前进""左转""提醒"关键词→force tool_choice\n'
        '4. 调用Kimi K2.5 API，返回多个tool_calls：[control_robot(forward,3s), control_robot(left), set_reminder("关火")]\n'
        '5. 命令队列按序执行：forward(3s)→wait→left→wait→HTTP POST /api/reminders\n'
        '6. 每个动作执行后通过WebSocket返回action feedback给前端，更新对话气泡状态\n\n'
        'B.2 提醒服务跨系统协同流程\n'
        'AI创建提醒: Web/鸿蒙→Python ai_chat_handler→Kimi tool_call→HTTP POST→Java /api/reminders→MySQL\n'
        'VOICE提醒发送: Java @Scheduled每分钟扫描→POST /internal/voice-reminder→Python→WS→Pi→Piper TTS→Speaker\n'
        'EMAIL提醒发送: Java @Scheduled每分钟扫描→EmailService→QQ SMTP→用户邮箱\n\n'
        'B.3 鸿蒙端WebSocket连接管理\n'
        'WebSocketService类实现了完整的连接生命周期管理：connect()→30秒内注册→30秒心跳→断开自动重连(最多5次/1.5s间隔)\n'
        '消息监听器模式(addMessageListener/removeMessageListener)支持多个页面组件同时订阅WS消息\n'
        '视频采用/video/frame端点轮询（鸿蒙不支持MJPEG流），通过Image组件加载JPEG快照实现视频预览'
    )
    add_formatted_paragraph(doc, tech_detail, font_size=Pt(9))

    # ========== SAVE ==========
    output_path = r'D:\yangy\桌面\01-作品说明文档_灵宠家伴_伏特加代码.docx'
    doc.save(output_path)
    print('Document saved successfully')
    return output_path

if __name__ == '__main__':
    create_document()
